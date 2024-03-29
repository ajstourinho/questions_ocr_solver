import json
import os
from docx import Document
from docx2pdf import convert
import glob
from docx.shared import Inches

def find_image_matching_json(json_name, search_directory):
    # Strip the .json extension to get the base name
    json_basename = os.path.splitext(json_name)[0]
    
    # Create a pattern to match files with the same base name as the JSON but with common image extensions
    pattern = os.path.join(search_directory, f"{json_basename}.*")
    
    # Use glob to find files matching the pattern
    for filename in glob.glob(pattern):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
            return filename  # Return the first matching image found
    
    return None  # No matching image found

def load_json_data(file_path):
    """Load and return JSON data from a file."""
    with open(file_path, 'r') as file:
        return json.load(file)

def create_pdf_from_docx(docx_path):
    """Convert a .docx file to PDF (may require admin permisions)."""
    convert(docx_path)

def add_data_docx(data, doc, file_name):
    """Add content from JSON data to a .docx file."""

    # Add original picture of questions
    image_filename = find_image_matching_json(file_name, "output_0_areas")
    doc.add_picture(image_filename, width=Inches(5.5))

    # Add content from JSON

    doc.add_paragraph().add_run(data['enunciado']).bold = True

    if data['tipo'] == "Discursiva":
        doc.add_paragraph(data['resposta'])
    if data['tipo'] == "Objetiva":
        for key, value in data['resposta'].items():
            if key != "alternativaCorreta":
                p1 = doc.add_paragraph()
                p1.add_run(key.upper() + ')  ' + value['alternativa']).bold = True
                
        doc.add_paragraph('\n')

        for key, value in data['resposta'].items():
            if key != "alternativaCorreta":
                p2 = doc.add_paragraph(key.upper() + ')  ' + value['alternativa'])
                
                p2.add_run(value['textoExplicativo'])

        doc.add_paragraph('\n')

        doc.add_paragraph('Alternatica correta: ' + data['resposta']['alternativaCorreta'])


if __name__ == "__main__":


    # Define JSON folder
    jsons_folder = 'output_1_jsons'

    # Create folder for the output files
    output_folder = 'output_2_docx_pdf'
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)


    # Create and docx
    docx_path = os.path.join(output_folder, 'questions.docx')
    doc = Document()
  
    # Iterate over each JSON file in the directory
    for file_name in os.listdir(jsons_folder):
        if file_name.endswith('.json'):
            file_path = os.path.join(jsons_folder, file_name)
            data = load_json_data(file_path)
            
            # Create docx
            add_data_docx(data, doc, file_name)

            # Add page break only if it is not the last one
            if file_name != os.listdir(jsons_folder)[-1]:
                doc.add_page_break()

    # Save .docx document
    doc.save(docx_path)
    print("\nThe docx file was saved.")

    # Check if there is Word installed to convert docx to pdf
    print("\nIn order to convert docx to pdf, with the docx2pdf library, it is required to have Word installed.\n")
    word_check = input("Is Word installed in this computer (y for Yes)? ")

    if word_check.lower() == "y":
        print("To convert docx to pdf, Word may open to ask for permission.")
        # Save the .pdf converted from .docx
        create_pdf_from_docx(docx_path)
    else:
        print("Unable to convert docx to pdf.")